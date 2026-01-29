from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from typing import Dict

class DocGenerator:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

    def generate_docx(self, template_name: str, data: Dict, output_name: str) -> str:
        """
        Génère un fichier DOCX structuré avec un style académique premium.
        """
        # Pour DocTerra, on génère un document structuré dynamiquement si aucun template n'est fourni
        doc = Document()
        
        # Styles de base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Page de Garde
        title = doc.add_heading(data.get('title', 'RAPPORT DE RECHERCHE DOCTERRA'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('\n' * 5)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Auteur : {data.get('author', 'Expert DocTerra')}")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph(f"Responsable : {data.get('professor', 'Core Agency')}")
        doc.add_paragraph(f"Date : {data.get('date', '2026')}")
        
        doc.add_page_break()

        # Sommaire
        doc.add_heading('Sommaire Exécutif', level=1)
        doc.add_paragraph(data.get('summary', 'Analyse automatisée des données extraites du web.'))
        
        doc.add_paragraph('\n')

        # Sections dynamiques
        sections = data.get('sections', [])
        for section in sections:
            doc.add_heading(section['title'], level=2)
            p = doc.add_paragraph(section['content'])
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            doc.add_paragraph('\n')

        # Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = "Généré par DocTerra Core v1.0.2 - Agence WebTerra"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        output_path = os.path.join(self.output_dir, f"{output_name}.docx")
        doc.save(output_path)
        return f"{output_name}.docx"
